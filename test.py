import time
import pyautogui
from docx import Document
from docx.shared import Pt
import keyboard
import pygetwindow as gw

def insert_formatted_text(formatted_string):
    # Create a new temporary document to handle the formatted text
    temp_doc = Document()
    parts = formatted_string.split('|')
    paragraph = temp_doc.add_paragraph()

    for i in range(0, len(parts), 3):  # Assuming 3 parts per set: text, formatting, size
        text = parts[i]
        formatting = parts[i + 1] if i + 1 < len(parts) else None
        size = parts[i + 2] if i + 2 < len(parts) else None

        run = paragraph.add_run(text)

        if formatting == 'bold':
            run.bold = True
        elif formatting == 'italic':
            run.italic = True

        if size:
            run.font.size = Pt(int(size.split(':')[1]))

    # Save the temporary document to a known location
    temp_doc.save('temp_document.docx')

    print("saved!")
    # Bring focus to the currently active window (the document in use)
    active_window = gw.getActiveWindow()
    if active_window:
        active_window.activate()
        time.sleep(1)  # Wait for the window to be active

    # Simulate Ctrl + O to open the temporary document
    pyautogui.hotkey('ctrl', 'o')
    time.sleep(1)  # Wait for the open dialog to appear

    # Type the name of the temporary document and hit Enter
    pyautogui.typewrite('temp_document.docx')
    pyautogui.press('enter')
    time.sleep(1)  # Wait for the document to open

    # Simulate Ctrl + A to select all text and then Ctrl + C to copy it
    pyautogui.hotkey('ctrl', 'a')
    pyautogui.hotkey('ctrl', 'c')

    # Switch back to the original document and paste the text
    active_window.activate()
    time.sleep(1)  # Wait for the window to be active
    pyautogui.hotkey('ctrl', 'v')  # Paste the copied text

    # Optionally, you could delete the temporary document afterward
    import os
    os.remove('temp_document.docx')

def insert_text_with_delay():
    print("Focus on your document window. You have 5 seconds...")
    time.sleep(5)  # Give the user time to focus on the document
    formatted_string = "This is bold|bold|size:12|This is italic|italic|size:16|Normal text|normal|size:10"
    insert_formatted_text(formatted_string)

# Monitor for the Ctrl key press
keyboard.add_hotkey('ctrl', insert_text_with_delay)

# Keep the program running
keyboard.wait('esc')  # Press 'esc' to exit
